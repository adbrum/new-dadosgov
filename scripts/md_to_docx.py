#!/usr/bin/env python3
"""Lightweight Markdown -> .docx converter (no pandoc needed).

Handles: ATX headings, paragraphs, bold (**), inline code (`), links, images,
bullet lists, blockquotes, fenced code blocks and GitHub pipe tables. Built for
the dados.gov.pt technical docs; not a full CommonMark implementation.

```mermaid fences are rendered to PNG with mermaid-cli (via npx) and embedded as
pictures; if mermaid-cli is unavailable the diagram source is kept as code text.
"""
import os
import re
import shutil
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
IMAGE_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")

# Widest picture that fits inside the default Word page margins.
PAGE_WIDTH_IN = 6.3


def add_runs(paragraph, text):
    """Render inline **bold**, *italic*, `code` and [text](url) spans."""
    text = LINK_RE.sub(r"\1", text)
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC0, 0x34, 0x21)
        else:
            paragraph.add_run(part)


def add_picture(doc, image_path, caption=""):
    """Insert a centred picture scaled to the text width, plus optional caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(PAGE_WIDTH_IN))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


def render_mermaid(source, out_path):
    """Render a mermaid diagram to PNG. Returns True on success."""
    mmdc = shutil.which("mmdc")
    cmd = [mmdc] if mmdc else ["npx", "-y", "@mermaid-js/mermaid-cli@11"]
    mmd_path = out_path.replace(".png", ".mmd")
    with open(mmd_path, "w", encoding="utf-8") as fh:
        fh.write(source)
    try:
        subprocess.run(
            cmd + ["-i", mmd_path, "-o", out_path, "-w", "2600", "-b", "white"],
            check=True,
            capture_output=True,
            timeout=600,
        )
        return os.path.exists(out_path)
    except (OSError, subprocess.SubprocessError):
        return False


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_separator(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def set_landscape(doc):
    """Switch the document to landscape A4 — useful for wide tables/diagrams."""
    global PAGE_WIDTH_IN
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
        section.left_margin = section.right_margin = Inches(0.6)
    PAGE_WIDTH_IN = 9.5


def convert(md_path, docx_path, landscape=False):
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    if landscape:
        set_landscape(doc)

    md_dir = os.path.dirname(os.path.abspath(md_path))
    tmp_dir = tempfile.mkdtemp(prefix="md_to_docx_")
    diagram_no = 0

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block
        if line.lstrip().startswith("```"):
            lang = line.lstrip()[3:].strip().lower()
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            source = "\n".join(buf)

            if lang == "mermaid":
                diagram_no += 1
                png = os.path.join(tmp_dir, f"diagram{diagram_no}.png")
                if render_mermaid(source, png):
                    add_picture(doc, png)
                    continue
                print(
                    f"WARN: mermaid-cli unavailable, diagram {diagram_no} kept as text",
                    file=sys.stderr,
                )

            p = doc.add_paragraph()
            run = p.add_run(source)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Image
        m = IMAGE_RE.match(line)
        if m:
            path = m.group(2).split(" ")[0]
            if not os.path.isabs(path):
                path = os.path.join(md_dir, path)
            if os.path.exists(path):
                add_picture(doc, path, m.group(1))
            else:
                print(f"WARN: image not found: {path}", file=sys.stderr)
            i += 1
            continue

        # Table (header row + separator)
        if line.strip().startswith("|") and i + 1 < n and is_separator(lines[i + 1]):
            header = split_row(line)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for idx, htext in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.paragraphs[0].clear()
                add_runs(cell.paragraphs[0], htext)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    text = row[idx] if idx < len(row) else ""
                    cells[idx].paragraphs[0].clear()
                    add_runs(cells[idx].paragraphs[0], text)
            doc.add_paragraph()
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading("", level=0)
                add_runs(h, text)
            else:
                h = doc.add_heading("", level=min(level - 1, 4))
                add_runs(h, text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            i += 1
            continue

        # Blockquote
        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, text)
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(2))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    argv = [a for a in sys.argv[1:] if a != "--landscape"]
    landscape = "--landscape" in sys.argv[1:]
    src = argv[0]
    dst = argv[1] if len(argv) > 1 else src.rsplit(".", 1)[0] + ".docx"
    convert(src, dst, landscape=landscape)
